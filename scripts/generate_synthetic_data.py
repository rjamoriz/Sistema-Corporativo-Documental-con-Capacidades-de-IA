"""
Generador de Datos Sintéticos para FinancIA 2030
Crea 200 documentos de prueba en diferentes formatos y categorías
"""
import os
import random
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict

from faker import Faker
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Inches, Pt
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont
import json

# Configurar Faker para español
fake = Faker(['es_ES', 'es_MX'])


class SyntheticDataGenerator:
    """Generador de documentos sintéticos"""
    
    def __init__(self, output_dir: str = "./synthetic_data"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Crear subdirectorios por categoría
        self.categories = {
            "legal": self.output_dir / "legal",
            "financial": self.output_dir / "financial",
            "hr": self.output_dir / "hr",
            "technical": self.output_dir / "technical",
            "marketing": self.output_dir / "marketing",
            "operations": self.output_dir / "operations",
            "compliance": self.output_dir / "compliance",
            "multimedia": self.output_dir / "multimedia"
        }
        
        for cat_dir in self.categories.values():
            cat_dir.mkdir(parents=True, exist_ok=True)
        
        print(f"📁 Output directory: {self.output_dir}")
    
    def generate_all(self, total_documents: int = 200):
        """Genera todos los documentos"""
        print(f"\n🚀 Starting generation of {total_documents} synthetic documents...")
        
        # Distribución por categoría
        distribution = {
            "legal": 30,      # Contratos, acuerdos
            "financial": 35,  # Facturas, presupuestos
            "hr": 25,         # Nóminas, contratos laborales
            "technical": 25,  # Manuales, especificaciones
            "marketing": 20,  # Campañas, informes
            "operations": 20, # Procedimientos, incidencias
            "compliance": 25, # Políticas, auditorías
            "multimedia": 20  # Imágenes, documentos escaneados
        }
        
        generated = 0
        
        # Generar por categoría
        print("\n" + "="*60)
        print("📄 LEGAL DOCUMENTS")
        print("="*60)
        for i in range(distribution["legal"]):
            self.generate_contract(i)
            generated += 1
            print(f"  ✅ Contract {i+1}/{distribution['legal']}")
        
        print("\n" + "="*60)
        print("💰 FINANCIAL DOCUMENTS")
        print("="*60)
        for i in range(distribution["financial"]):
            if i % 2 == 0:
                self.generate_invoice(i)
                print(f"  ✅ Invoice {i+1}")
            else:
                self.generate_budget(i)
                print(f"  ✅ Budget {i+1}")
            generated += 1
        
        print("\n" + "="*60)
        print("👥 HR DOCUMENTS")
        print("="*60)
        for i in range(distribution["hr"]):
            if i % 2 == 0:
                self.generate_payroll(i)
                print(f"  ✅ Payroll {i+1}")
            else:
                self.generate_employment_contract(i)
                print(f"  ✅ Employment contract {i+1}")
            generated += 1
        
        print("\n" + "="*60)
        print("⚙️ TECHNICAL DOCUMENTS")
        print("="*60)
        for i in range(distribution["technical"]):
            self.generate_technical_spec(i)
            generated += 1
            print(f"  ✅ Technical spec {i+1}/{distribution['technical']}")
        
        print("\n" + "="*60)
        print("📊 MARKETING DOCUMENTS")
        print("="*60)
        for i in range(distribution["marketing"]):
            self.generate_marketing_report(i)
            generated += 1
            print(f"  ✅ Marketing report {i+1}/{distribution['marketing']}")
        
        print("\n" + "="*60)
        print("🔧 OPERATIONAL DOCUMENTS")
        print("="*60)
        for i in range(distribution["operations"]):
            self.generate_operational_doc(i)
            generated += 1
            print(f"  ✅ Operational doc {i+1}/{distribution['operations']}")
        
        print("\n" + "="*60)
        print("📋 COMPLIANCE DOCUMENTS")
        print("="*60)
        for i in range(distribution["compliance"]):
            self.generate_compliance_policy(i)
            generated += 1
            print(f"  ✅ Compliance policy {i+1}/{distribution['compliance']}")
        
        print("\n" + "="*60)
        print("🖼️ MULTIMEDIA DOCUMENTS")
        print("="*60)
        for i in range(distribution["multimedia"]):
            if i % 2 == 0:
                self.generate_scanned_document(i)
                print(f"  ✅ Scanned doc {i+1}")
            else:
                self.generate_infographic(i)
                print(f"  ✅ Infographic {i+1}")
            generated += 1
        
        print("\n" + "="*60)
        print(f"✅ GENERATION COMPLETE!")
        print(f"📊 Total documents generated: {generated}")
        print(f"📁 Location: {self.output_dir}")
        print("="*60 + "\n")
        
        # Generar manifiesto
        self._generate_manifest(generated)
    
    def generate_contract(self, index: int):
        """Genera un contrato legal en PDF"""
        filename = self.categories["legal"] / f"contract_{index:03d}.pdf"
        
        c = canvas.Canvas(str(filename), pagesize=A4)
        width, height = A4
        
        # Título
        c.setFont("Helvetica-Bold", 16)
        c.drawString(2*cm, height - 3*cm, "CONTRATO DE PRESTACIÓN DE SERVICIOS")
        
        # Partes
        c.setFont("Helvetica", 11)
        y = height - 5*cm
        
        company1 = fake.company()
        company2 = fake.company()
        cif1 = fake.bothify(text='?########?').upper()
        cif2 = fake.bothify(text='?########?').upper()
        
        c.drawString(2*cm, y, f"REUNIDOS")
        y -= 0.7*cm
        c.drawString(2*cm, y, f"De una parte, {company1}, con CIF {cif1}, representada por")
        y -= 0.5*cm
        c.drawString(2*cm, y, f"{fake.name()}, con DNI {fake.bothify(text='########?').upper()}")
        y -= 0.7*cm
        c.drawString(2*cm, y, f"De otra parte, {company2}, con CIF {cif2}, representada por")
        y -= 0.5*cm
        c.drawString(2*cm, y, f"{fake.name()}, con DNI {fake.bothify(text='########?').upper()}")
        
        y -= 1.5*cm
        c.drawString(2*cm, y, "CLÁUSULAS")
        y -= 0.7*cm
        
        # Cláusulas
        clauses = [
            f"PRIMERA. - Objeto del contrato: {company1} se compromete a prestar servicios de {fake.job()} a {company2}.",
            f"SEGUNDA. - Duración: El presente contrato tendrá una duración de {random.randint(12, 36)} meses desde su firma.",
            f"TERCERA. - Precio: El precio total asciende a {random.randint(10000, 100000):,}€ más IVA.",
            f"CUARTA. - Forma de pago: Pagos mensuales de {random.randint(500, 5000):,}€ mediante transferencia bancaria.",
            "QUINTA. - Jurisdicción: Para cualquier litigio, las partes se someten a los juzgados de Madrid.",
        ]
        
        for i, clause in enumerate(clauses):
            c.drawString(2*cm, y, clause)
            y -= 0.7*cm
            if y < 3*cm:
                c.showPage()
                y = height - 3*cm
        
        # Firma
        y -= 1*cm
        date = fake.date_between(start_date='-2y', end_date='today')
        c.drawString(2*cm, y, f"Firmado en Madrid, a {date.strftime('%d de %B de %Y')}")
        y -= 2*cm
        c.drawString(2*cm, y, f"_____________________")
        c.drawString(12*cm, y, f"_____________________")
        y -= 0.5*cm
        c.drawString(2*cm, y, f"{company1}")
        c.drawString(12*cm, y, f"{company2}")
        
        c.save()
    
    def generate_invoice(self, index: int):
        """Genera una factura en PDF"""
        filename = self.categories["financial"] / f"invoice_{index:03d}.pdf"
        
        c = canvas.Canvas(str(filename), pagesize=A4)
        width, height = A4
        
        # Cabecera
        c.setFont("Helvetica-Bold", 18)
        c.drawString(2*cm, height - 3*cm, "FACTURA")
        
        invoice_number = f"F-{fake.year()}-{index:04d}"
        c.setFont("Helvetica", 11)
        c.drawString(2*cm, height - 4*cm, f"Número: {invoice_number}")
        c.drawString(2*cm, height - 4.5*cm, f"Fecha: {fake.date_this_year()}")
        
        # Emisor
        y = height - 6*cm
        company = fake.company()
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, y, "EMISOR:")
        y -= 0.7*cm
        c.setFont("Helvetica", 10)
        c.drawString(2*cm, y, company)
        y -= 0.5*cm
        c.drawString(2*cm, y, f"CIF: {fake.bothify(text='?########?').upper()}")
        y -= 0.5*cm
        c.drawString(2*cm, y, fake.address())
        
        # Cliente
        y -= 1.5*cm
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, y, "CLIENTE:")
        y -= 0.7*cm
        c.setFont("Helvetica", 10)
        c.drawString(2*cm, y, fake.company())
        y -= 0.5*cm
        c.drawString(2*cm, y, f"CIF: {fake.bothify(text='?########?').upper()}")
        
        # Líneas de factura
        y -= 2*cm
        c.setFont("Helvetica-Bold", 11)
        c.drawString(2*cm, y, "Concepto")
        c.drawString(12*cm, y, "Cantidad")
        c.drawString(15*cm, y, "Precio")
        c.drawString(18*cm, y, "Total")
        
        y -= 0.5*cm
        c.line(2*cm, y, 19*cm, y)
        
        y -= 0.7*cm
        c.setFont("Helvetica", 10)
        
        subtotal = 0
        num_items = random.randint(3, 8)
        for i in range(num_items):
            concept = fake.catch_phrase()
            quantity = random.randint(1, 10)
            price = random.randint(50, 500)
            total = quantity * price
            subtotal += total
            
            c.drawString(2*cm, y, concept[:30])
            c.drawString(12*cm, y, str(quantity))
            c.drawString(15*cm, y, f"{price}€")
            c.drawString(18*cm, y, f"{total}€")
            y -= 0.6*cm
        
        # Totales
        y -= 1*cm
        iva = subtotal * 0.21
        total = subtotal + iva
        
        c.setFont("Helvetica-Bold", 11)
        c.drawString(15*cm, y, "Subtotal:")
        c.drawString(18*cm, y, f"{subtotal:.2f}€")
        y -= 0.6*cm
        c.drawString(15*cm, y, "IVA (21%):")
        c.drawString(18*cm, y, f"{iva:.2f}€")
        y -= 0.6*cm
        c.setFont("Helvetica-Bold", 13)
        c.drawString(15*cm, y, "TOTAL:")
        c.drawString(18*cm, y, f"{total:.2f}€")
        
        c.save()
    
    def generate_budget(self, index: int):
        """Genera un presupuesto en Excel"""
        filename = self.categories["financial"] / f"budget_{index:03d}.xlsx"
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Presupuesto"
        
        # Cabecera
        ws['A1'] = "PRESUPUESTO"
        ws['A2'] = f"Número: PRE-{fake.year()}-{index:04d}"
        ws['A3'] = f"Fecha: {fake.date_this_year()}"
        ws['A4'] = f"Empresa: {fake.company()}"
        ws['A5'] = f"Cliente: {fake.company()}"
        
        # Headers
        headers = ['Concepto', 'Cantidad', 'Precio Unitario', 'Total']
        for col, header in enumerate(headers, start=1):
            ws.cell(row=7, column=col, value=header)
        
        # Líneas
        row = 8
        subtotal = 0
        for i in range(random.randint(5, 15)):
            concept = fake.catch_phrase()
            quantity = random.randint(1, 20)
            price = random.randint(10, 1000)
            total = quantity * price
            subtotal += total
            
            ws.cell(row=row, column=1, value=concept)
            ws.cell(row=row, column=2, value=quantity)
            ws.cell(row=row, column=3, value=price)
            ws.cell(row=row, column=4, value=total)
            row += 1
        
        # Totales
        row += 1
        ws.cell(row=row, column=3, value="Subtotal:")
        ws.cell(row=row, column=4, value=subtotal)
        row += 1
        iva = subtotal * 0.21
        ws.cell(row=row, column=3, value="IVA (21%):")
        ws.cell(row=row, column=4, value=round(iva, 2))
        row += 1
        ws.cell(row=row, column=3, value="TOTAL:")
        ws.cell(row=row, column=4, value=round(subtotal + iva, 2))
        
        wb.save(filename)
    
    def generate_payroll(self, index: int):
        """Genera una nómina en PDF"""
        filename = self.categories["hr"] / f"payroll_{index:03d}.pdf"
        
        c = canvas.Canvas(str(filename), pagesize=A4)
        width, height = A4
        
        # Título
        c.setFont("Helvetica-Bold", 16)
        c.drawString(2*cm, height - 3*cm, "NÓMINA")
        
        # Datos
        c.setFont("Helvetica", 11)
        y = height - 5*cm
        
        employee = fake.name()
        dni = fake.bothify(text='########?').upper()
        company = fake.company()
        
        month = fake.month_name()
        year = fake.year()
        
        c.drawString(2*cm, y, f"Trabajador: {employee}")
        y -= 0.6*cm
        c.drawString(2*cm, y, f"DNI: {dni}")
        y -= 0.6*cm
        c.drawString(2*cm, y, f"Empresa: {company}")
        y -= 0.6*cm
        c.drawString(2*cm, y, f"Período: {month} {year}")
        
        # Devengos
        y -= 1.5*cm
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, y, "DEVENGOS")
        y -= 0.8*cm
        c.setFont("Helvetica", 10)
        
        base_salary = random.randint(1200, 3000)
        c.drawString(2*cm, y, f"Salario base: {base_salary}€")
        y -= 0.6*cm
        
        complement = random.randint(100, 500)
        c.drawString(2*cm, y, f"Complementos: {complement}€")
        y -= 0.6*cm
        
        total_devengos = base_salary + complement
        c.setFont("Helvetica-Bold", 11)
        c.drawString(2*cm, y, f"Total devengos: {total_devengos}€")
        
        # Deducciones
        y -= 1.5*cm
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, y, "DEDUCCIONES")
        y -= 0.8*cm
        c.setFont("Helvetica", 10)
        
        ss = round(total_devengos * 0.0635, 2)
        c.drawString(2*cm, y, f"Seguridad Social (6.35%): {ss}€")
        y -= 0.6*cm
        
        irpf = round(total_devengos * 0.15, 2)
        c.drawString(2*cm, y, f"IRPF (15%): {irpf}€")
        y -= 0.6*cm
        
        total_deducciones = ss + irpf
        c.setFont("Helvetica-Bold", 11)
        c.drawString(2*cm, y, f"Total deducciones: {total_deducciones}€")
        
        # Líquido
        y -= 1.5*cm
        liquido = total_devengos - total_deducciones
        c.setFont("Helvetica-Bold", 14)
        c.drawString(2*cm, y, f"LÍQUIDO A PERCIBIR: {liquido:.2f}€")
        
        c.save()
    
    def generate_employment_contract(self, index: int):
        """Genera un contrato laboral en DOCX"""
        filename = self.categories["hr"] / f"employment_contract_{index:03d}.docx"
        
        doc = Document()
        
        # Título
        title = doc.add_heading('CONTRATO DE TRABAJO', level=1)
        title.alignment = 1  # Center
        
        # Datos
        doc.add_paragraph(f"Fecha: {fake.date_this_year()}")
        doc.add_paragraph()
        
        company = fake.company()
        employee = fake.name()
        dni = fake.bothify(text='########?').upper()
        position = fake.job()
        salary = random.randint(18000, 45000)
        
        doc.add_heading('PARTES', level=2)
        doc.add_paragraph(f"EMPRESA: {company}")
        doc.add_paragraph(f"CIF: {fake.bothify(text='?########?').upper()}")
        doc.add_paragraph(f"Dirección: {fake.address()}")
        doc.add_paragraph()
        doc.add_paragraph(f"TRABAJADOR: {employee}")
        doc.add_paragraph(f"DNI: {dni}")
        doc.add_paragraph(f"Domicilio: {fake.address()}")
        doc.add_paragraph()
        
        doc.add_heading('CLÁUSULAS', level=2)
        
        doc.add_heading('PRIMERA. Puesto de trabajo', level=3)
        doc.add_paragraph(
            f"El trabajador prestará sus servicios como {position} en las instalaciones "
            f"de la empresa, desempeñando las funciones propias del cargo."
        )
        
        doc.add_heading('SEGUNDA. Duración', level=3)
        doc.add_paragraph(
            f"El presente contrato tiene carácter indefinido, con un período de prueba "
            f"de {random.choice([1, 2, 3, 6])} meses."
        )
        
        doc.add_heading('TERCERA. Retribución', level=3)
        doc.add_paragraph(
            f"El trabajador percibirá un salario bruto anual de {salary:,}€, "
            f"distribuido en {random.choice([12, 14])} pagas."
        )
        
        doc.add_heading('CUARTA. Jornada', level=3)
        doc.add_paragraph(
            f"La jornada laboral será de {random.choice([35, 37.5, 40])} horas semanales, "
            f"de lunes a viernes."
        )
        
        doc.add_heading('QUINTA. Vacaciones', level=3)
        doc.add_paragraph(
            "El trabajador tiene derecho a 30 días naturales de vacaciones anuales retribuidas."
        )
        
        doc.save(filename)
    
    def generate_technical_spec(self, index: int):
        """Genera una especificación técnica en DOCX"""
        filename = self.categories["technical"] / f"tech_spec_{index:03d}.docx"
        
        doc = Document()
        
        # Título
        title = doc.add_heading('ESPECIFICACIÓN TÉCNICA', level=1)
        
        system_name = f"Sistema {fake.word().capitalize()} {random.randint(1000, 9999)}"
        doc.add_heading(system_name, level=2)
        
        doc.add_paragraph(f"Versión: {random.randint(1, 5)}.{random.randint(0, 9)}.{random.randint(0, 20)}")
        doc.add_paragraph(f"Fecha: {fake.date_this_year()}")
        doc.add_paragraph(f"Autor: {fake.name()}")
        doc.add_paragraph()
        
        doc.add_heading('1. Introducción', level=2)
        doc.add_paragraph(
            f"Este documento describe la arquitectura y funcionamiento del {system_name}. "
            f"El sistema tiene como objetivo proporcionar {fake.catch_phrase()}."
        )
        
        doc.add_heading('2. Arquitectura', level=2)
        doc.add_paragraph("El sistema se compone de los siguientes módulos:")
        
        for i in range(random.randint(3, 6)):
            doc.add_paragraph(
                f"• Módulo {fake.word().capitalize()}: Responsable de {fake.bs()}",
                style='List Bullet'
            )
        
        doc.add_heading('3. Requisitos', level=2)
        doc.add_paragraph("Requisitos funcionales:")
        for i in range(random.randint(4, 8)):
            doc.add_paragraph(
                f"RF-{i+1}: El sistema debe {fake.bs()}",
                style='List Bullet'
            )
        
        doc.add_paragraph()
        doc.add_paragraph("Requisitos no funcionales:")
        doc.add_paragraph("• Disponibilidad: 99.9%", style='List Bullet')
        doc.add_paragraph(f"• Tiempo de respuesta: < {random.randint(100, 500)}ms", style='List Bullet')
        doc.add_paragraph(f"• Usuarios concurrentes: {random.randint(100, 10000)}", style='List Bullet')
        
        doc.add_heading('4. Tecnologías', level=2)
        techs = random.sample(['Python', 'Java', 'Node.js', 'React', 'Angular', 'Vue.js', 
                               'PostgreSQL', 'MongoDB', 'Redis', 'Kafka', 'Docker', 'Kubernetes'], 
                              k=random.randint(4, 7))
        for tech in techs:
            doc.add_paragraph(f"• {tech}", style='List Bullet')
        
        doc.save(filename)
    
    def generate_marketing_report(self, index: int):
        """Genera un informe de marketing en PDF"""
        filename = self.categories["marketing"] / f"marketing_report_{index:03d}.pdf"
        
        c = canvas.Canvas(str(filename), pagesize=A4)
        width, height = A4
        
        c.setFont("Helvetica-Bold", 18)
        c.drawString(2*cm, height - 3*cm, "INFORME DE CAMPAÑA DE MARKETING")
        
        c.setFont("Helvetica", 11)
        y = height - 5*cm
        
        campaign = fake.catch_phrase()
        c.drawString(2*cm, y, f"Campaña: {campaign}")
        y -= 0.6*cm
        c.drawString(2*cm, y, f"Período: Q{random.randint(1, 4)} {fake.year()}")
        y -= 0.6*cm
        c.drawString(2*cm, y, f"Responsable: {fake.name()}")
        
        y -= 1.5*cm
        c.setFont("Helvetica-Bold", 13)
        c.drawString(2*cm, y, "MÉTRICAS CLAVE")
        
        y -= 1*cm
        c.setFont("Helvetica", 11)
        
        metrics = [
            ("Inversión total", f"{random.randint(5000, 50000):,}€"),
            ("Impresiones", f"{random.randint(100000, 1000000):,}"),
            ("Clicks", f"{random.randint(5000, 50000):,}"),
            ("CTR", f"{random.uniform(1, 5):.2f}%"),
            ("Conversiones", f"{random.randint(500, 5000):,}"),
            ("Coste por conversión", f"{random.uniform(5, 50):.2f}€"),
            ("ROI", f"{random.uniform(150, 300):.1f}%"),
        ]
        
        for metric, value in metrics:
            c.drawString(2*cm, y, f"{metric}:")
            c.drawString(10*cm, y, value)
            y -= 0.7*cm
        
        y -= 1*cm
        c.setFont("Helvetica-Bold", 13)
        c.drawString(2*cm, y, "CONCLUSIONES")
        y -= 0.8*cm
        c.setFont("Helvetica", 10)
        
        conclusions = [
            f"La campaña ha superado las expectativas iniciales en un {random.randint(10, 50)}%.",
            f"El canal con mejor rendimiento ha sido {random.choice(['Google Ads', 'Facebook', 'Instagram', 'LinkedIn'])}.",
            f"Se recomienda aumentar la inversión en {random.choice(['contenido', 'paid media', 'influencers'])} para el próximo trimestre.",
        ]
        
        for conclusion in conclusions:
            c.drawString(2*cm, y, conclusion)
            y -= 0.7*cm
        
        c.save()
    
    def generate_operational_doc(self, index: int):
        """Genera un documento operacional"""
        filename = self.categories["operations"] / f"operational_{index:03d}.docx"
        
        doc = Document()
        
        title = doc.add_heading('PROCEDIMIENTO OPERACIONAL', level=1)
        
        proc_name = f"PO-{index:03d}: {fake.catch_phrase()}"
        doc.add_heading(proc_name, level=2)
        
        doc.add_paragraph(f"Versión: 1.{random.randint(0, 5)}")
        doc.add_paragraph(f"Fecha de aprobación: {fake.date_this_year()}")
        doc.add_paragraph(f"Responsable: {fake.name()}")
        doc.add_paragraph()
        
        doc.add_heading('1. Objetivo', level=2)
        doc.add_paragraph(
            f"Establecer el procedimiento para {fake.bs()} de manera eficiente y segura."
        )
        
        doc.add_heading('2. Alcance', level=2)
        doc.add_paragraph(
            f"Este procedimiento aplica a {random.choice(['todo el personal', 'el departamento técnico', 'los supervisores'])} "
            f"de {fake.company()}."
        )
        
        doc.add_heading('3. Responsabilidades', level=2)
        roles = ['Gerente', 'Supervisor', 'Técnico', 'Operario']
        for role in roles:
            doc.add_paragraph(
                f"• {role}: {fake.bs()}",
                style='List Bullet'
            )
        
        doc.add_heading('4. Procedimiento', level=2)
        for i in range(random.randint(5, 10)):
            doc.add_paragraph(
                f"Paso {i+1}: {fake.sentence()}",
                style='List Number'
            )
        
        doc.add_heading('5. Registros', level=2)
        doc.add_paragraph(
            f"Los registros generados deben conservarse durante {random.randint(1, 5)} años."
        )
        
        doc.save(filename)
    
    def generate_compliance_policy(self, index: int):
        """Genera una política de cumplimiento"""
        filename = self.categories["compliance"] / f"policy_{index:03d}.pdf"
        
        c = canvas.Canvas(str(filename), pagesize=A4)
        width, height = A4
        
        c.setFont("Helvetica-Bold", 16)
        policy_type = random.choice([
            "Política de Protección de Datos",
            "Política de Seguridad de la Información",
            "Política de Cumplimiento GDPR",
            "Política de Gestión de Riesgos",
            "Política de Ética y Conducta"
        ])
        c.drawString(2*cm, height - 3*cm, policy_type)
        
        c.setFont("Helvetica", 11)
        y = height - 5*cm
        
        c.drawString(2*cm, y, f"Empresa: {fake.company()}")
        y -= 0.6*cm
        c.drawString(2*cm, y, f"Versión: {random.randint(1, 5)}.0")
        y -= 0.6*cm
        c.drawString(2*cm, y, f"Fecha de aprobación: {fake.date_this_year()}")
        y -= 0.6*cm
        c.drawString(2*cm, y, f"Aprobado por: {fake.name()} (DPO)")
        
        y -= 1.5*cm
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, y, "1. OBJETO")
        y -= 0.8*cm
        c.setFont("Helvetica", 10)
        c.drawString(2*cm, y, "Establecer los principios y directrices para garantizar el cumplimiento de")
        y -= 0.5*cm
        c.drawString(2*cm, y, "las normativas vigentes en materia de protección de datos y privacidad.")
        
        y -= 1*cm
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, y, "2. ALCANCE")
        y -= 0.8*cm
        c.setFont("Helvetica", 10)
        c.drawString(2*cm, y, "Esta política aplica a todos los empleados, colaboradores y terceros que")
        y -= 0.5*cm
        c.drawString(2*cm, y, "tengan acceso a datos personales bajo responsabilidad de la organización.")
        
        y -= 1*cm
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, y, "3. PRINCIPIOS")
        y -= 0.8*cm
        c.setFont("Helvetica", 10)
        
        principles = [
            "Licitud, lealtad y transparencia en el tratamiento",
            "Limitación de la finalidad",
            "Minimización de datos",
            "Exactitud",
            "Limitación del plazo de conservación",
            "Integridad y confidencialidad",
            "Responsabilidad proactiva"
        ]
        
        for principle in principles:
            c.drawString(2.5*cm, y, f"• {principle}")
            y -= 0.6*cm
            if y < 3*cm:
                c.showPage()
                y = height - 3*cm
        
        c.save()
    
    def generate_scanned_document(self, index: int):
        """Genera un documento simulando escaneo (imagen con texto)"""
        filename = self.categories["multimedia"] / f"scanned_{index:03d}.png"
        
        # Crear imagen
        img = Image.new('RGB', (2480, 3508), color='white')  # A4 at 300 DPI
        draw = ImageDraw.Draw(img)
        
        # Simular textura de papel escaneado
        import numpy as np
        arr = np.array(img)
        noise = np.random.randint(-10, 10, arr.shape, dtype='int16')
        arr = np.clip(arr + noise, 0, 255).astype('uint8')
        img = Image.fromarray(arr)
        draw = ImageDraw.Draw(img)
        
        # Añadir texto
        try:
            from PIL import ImageFont
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 40)
        except:
            font = ImageFont.load_default()
        
        y = 200
        draw.text((200, y), fake.company(), fill='black', font=font)
        y += 100
        draw.text((200, y), f"Documento: DOC-{index:04d}", fill='black', font=font)
        y += 80
        draw.text((200, y), f"Fecha: {fake.date_this_year()}", fill='black', font=font)
        y += 120
        
        # Añadir párrafos de texto
        for _ in range(15):
            text = fake.sentence()
            draw.text((200, y), text, fill='black', font=font)
            y += 60
        
        # Simular ligera inclinación del escaneo
        img = img.rotate(random.uniform(-0.5, 0.5), fillcolor='white')
        
        img.save(filename, 'PNG')
    
    def generate_infographic(self, index: int):
        """Genera una infografía simple"""
        filename = self.categories["multimedia"] / f"infographic_{index:03d}.png"
        
        # Crear imagen
        img = Image.new('RGB', (1200, 1600), color='#f0f0f0')
        draw = ImageDraw.Draw(img)
        
        # Título
        colors = ['#3498db', '#e74c3c', '#2ecc71', '#f39c12', '#9b59b6']
        color = random.choice(colors)
        
        draw.rectangle([0, 0, 1200, 200], fill=color)
        
        try:
            font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 60)
            font_text = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 30)
        except:
            font_title = ImageFont.load_default()
            font_text = ImageFont.load_default()
        
        draw.text((50, 70), fake.catch_phrase()[:30], fill='white', font=font_title)
        
        # Bloques de información
        y = 250
        for i in range(4):
            block_color = random.choice(colors)
            draw.rectangle([100, y, 1100, y+250], fill='white', outline=block_color, width=5)
            draw.text((150, y+50), f"Métrica {i+1}", fill=block_color, font=font_title)
            draw.text((150, y+130), f"{random.randint(100, 999)}", fill='black', font=font_title)
            draw.text((150, y+190), fake.sentence()[:40], fill='#666', font=font_text)
            y += 300
        
        img.save(filename, 'PNG')
    
    def _generate_manifest(self, total_generated: int):
        """Genera un archivo manifest con información de los documentos generados"""
        manifest = {
            "generation_date": datetime.now().isoformat(),
            "total_documents": total_generated,
            "categories": {},
            "statistics": {
                "formats": {},
                "total_size_mb": 0
            }
        }
        
        # Contar documentos por categoría y formato
        for category, cat_dir in self.categories.items():
            files = list(cat_dir.glob("*"))
            manifest["categories"][category] = {
                "count": len(files),
                "files": [f.name for f in files]
            }
            
            for file in files:
                ext = file.suffix[1:]  # Remove dot
                manifest["statistics"]["formats"][ext] = manifest["statistics"]["formats"].get(ext, 0) + 1
                manifest["statistics"]["total_size_mb"] += file.stat().st_size / (1024 * 1024)
        
        manifest["statistics"]["total_size_mb"] = round(manifest["statistics"]["total_size_mb"], 2)
        
        # Guardar manifest
        manifest_file = self.output_dir / "manifest.json"
        with open(manifest_file, 'w', encoding='utf-8') as f:
            json.dump(manifest, f, indent=2, ensure_ascii=False)
        
        print(f"\n📋 Manifest saved: {manifest_file}")
        print(f"📊 Total size: {manifest['statistics']['total_size_mb']} MB")
        print(f"📁 Formats: {manifest['statistics']['formats']}")


def main():
    """Función principal"""
    print("\n" + "="*70)
    print("  FinancIA 2030 - Synthetic Data Generator")
    print("="*70)
    
    generator = SyntheticDataGenerator()
    generator.generate_all(total_documents=200)
    
    print("\n🎉 All done! You can now upload these documents to test the system.")
    print("="*70 + "\n")


if __name__ == "__main__":
    main()
